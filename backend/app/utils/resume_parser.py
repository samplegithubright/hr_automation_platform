import io
import os
import docx
import httpx
from pypdf import PdfReader

class ResumeParser:
    @classmethod
    async def extract_text(cls, file_path_or_url: str) -> str:
        """
        Extracts plain text from a PDF or DOCX file.
        Works with local file paths and remote S3 URLs.
        """
        # If it's an HTTP/HTTPS URL, download it into memory
        if file_path_or_url.startswith("http://") or file_path_or_url.startswith("https://"):
            async with httpx.AsyncClient() as client:
                response = await client.get(file_path_or_url)
                if response.status_code != 200:
                    raise Exception(f"Failed to download resume from S3 URL: {file_path_or_url}")
                file_bytes = response.content
                filename = file_path_or_url.split("/")[-1]
        else:
            # Otherwise, read from local disk
            if not os.path.exists(file_path_or_url):
                raise FileNotFoundError(f"Local resume file not found: {file_path_or_url}")
            with open(file_path_or_url, "rb") as f:
                file_bytes = f.read()
                filename = os.path.basename(file_path_or_url)

        # Parse based on extension / filename
        lower_name = filename.lower()
        if lower_name.endswith(".pdf"):
            return cls._parse_pdf(file_bytes)
        elif lower_name.endswith(".docx"):
            return cls._parse_docx(file_bytes)
        else:
            # Fallback check of magic bytes if extension is missing/weird
            if file_bytes.startswith(b'%PDF'):
                return cls._parse_pdf(file_bytes)
            elif file_bytes.startswith(b'PK\x03\x04'):
                return cls._parse_docx(file_bytes)
            raise ValueError("Unsupported resume file format. Only PDF and DOCX are supported.")

    @staticmethod
    def _parse_pdf(file_bytes: bytes) -> str:
        """Parses a PDF from bytes in-memory and extracts text."""
        pdf_file = io.BytesIO(file_bytes)
        reader = PdfReader(pdf_file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()

    @staticmethod
    def _parse_docx(file_bytes: bytes) -> str:
        """Parses a DOCX from bytes in-memory and extracts text."""
        docx_file = io.BytesIO(file_bytes)
        doc = docx.Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
        # Also parse tables if present (optional, but highly recommended for resumes)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        text += cell.text + " "
                text += "\n"
        return text.strip()
